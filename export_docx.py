"""
Export markdown content to a .docx file.

Usage:
    python export_docx.py --output <path/to/file.docx> < content.md
"""

import argparse
import re
import sys
from pathlib import Path


def md_to_docx(md_text: str, output_path: Path) -> None:
    from docx import Document

    doc = Document()
    lines = md_text.split('\n')
    i = 0

    def strip_inline(text: str) -> str:
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*',     r'\1', text)
        text = re.sub(r'`(.+?)`',       r'\1', text)
        return text

    while i < len(lines):
        s = lines[i].strip()

        if s.startswith('### '):
            doc.add_heading(strip_inline(s[4:]), level=3)
            i += 1
        elif s.startswith('## '):
            doc.add_heading(strip_inline(s[3:]), level=2)
            i += 1
        elif s.startswith('# '):
            doc.add_heading(strip_inline(s[2:]), level=1)
            i += 1
        elif s.startswith('|'):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            data_rows = [r for r in table_lines
                         if not re.match(r'^\|[-:| ]+\|$', r.replace(' ', ''))]
            if data_rows:
                ncols = len([c.strip() for c in data_rows[0].strip('|').split('|')])
                tbl = doc.add_table(rows=len(data_rows), cols=ncols)
                tbl.style = 'Table Grid'
                for ri, row_str in enumerate(data_rows):
                    cells = [c.strip() for c in row_str.strip('|').split('|')]
                    for ci in range(ncols):
                        cell_text = strip_inline(cells[ci]) if ci < len(cells) else ''
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = cell_text
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
        elif s.startswith('- ') or s.startswith('* '):
            doc.add_paragraph(strip_inline(s[2:]), style='List Bullet')
            i += 1
        elif s in ('', '---'):
            i += 1
        else:
            if s:
                doc.add_paragraph(strip_inline(s))
            i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert markdown stdin to .docx')
    parser.add_argument('--output', required=True, help='Output .docx file path')
    args = parser.parse_args()

    md_text = sys.stdin.read()
    md_to_docx(md_text, Path(args.output))
    print('OK')
