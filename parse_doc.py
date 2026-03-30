"""
DOC/DOCX parser for RAG compliance system.

Parses all Word documents found in List, Procedure, Datasheet, and Specification
folders across engineering disciplines.

Two parsing modes:
  - PRC / SPC files  → section-heading chunking (split on Heading 1/2 styles)
  - LST .doc files   → table extraction (row-by-row, same approach as XLS parser)

Legacy .doc files are converted to .docx via LibreOffice before parsing.

Usage:
    python parse_doc.py <root_dir>
    python parse_doc.py .

Dependencies:
    pip install python-docx
    LibreOffice must be installed (soffice in PATH or /opt/homebrew/bin/soffice)
"""

import re
import json
import subprocess
import tempfile
import shutil
from pathlib import Path

import docx


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

ALLOWED_FOLDERS = {'List', 'Procedure', 'Procedures'}
# Datasheet and Specification excluded for now — may be added in future

HEADING_STYLES = {'Heading 1', 'Heading 2', 'Heading 3',
                  'heading 1', 'heading 2', 'heading 3'}

CONSTRAINT_PATTERN = re.compile(
    r'\b(must|shall|required|mandatory|not permitted|shall not)\b|[<>≤≥]=?\s*\d',
    re.IGNORECASE
)

SOFFICE = next(
    (p for p in ['/opt/homebrew/bin/soffice', '/usr/bin/soffice', 'soffice']
     if shutil.which(p) or Path(p).exists()),
    'soffice'
)


# ---------------------------------------------------------------------------
# Filename metadata detection (shared logic with parse_xls.py)
# ---------------------------------------------------------------------------

def detect_discipline(filename: str) -> int | None:
    match = re.match(r'^(\d+)-', filename)
    return int(match.group(1)) if match else None

def detect_doc_id(filename: str) -> str:
    match = re.match(r'^(\d+-[A-Z]+-[\d\-]+)', filename)
    return match.group(1).rstrip('-') if match else filename

def detect_doc_type(filename: str) -> str:
    for t in ('LST', 'PRC', 'DST', 'SPC', 'FRM', 'REP'):
        if f'-{t}-' in filename:
            return t
    return 'UNKNOWN'

def detect_revision(filename: str) -> str | None:
    match = re.search(r'-R(\d+)', filename, re.IGNORECASE)
    return f'R{match.group(1)}' if match else None

def detect_revision_number(filename: str) -> int:
    match = re.search(r'-R(\d+)', filename, re.IGNORECASE)
    return int(match.group(1)) if match else -1

def detect_is_template(filename: str) -> bool:
    return bool(re.search(r'\[?[Tt]emplate\]?|\([Tt]emplate\)', filename))

def detect_is_hold(filename: str) -> bool:
    return 'HOLD' in filename.upper() and 'DO NOT USE' in filename.upper()

DISCIPLINE_NAMES = {
    0: 'Administration',
    1: 'Process Technology',
    2: 'Civil & Structural',
    3: 'Mechanical',
    4: 'Equipment',
    5: 'Piping & Layout',
    6: 'Insulation',
    7: 'Coatings',
    8: 'Instrumentation & Control',
    9: 'Electrical',
}

def detect_source_folder(file_path: Path) -> str:
    for part in reversed(file_path.parts):
        if part in ('Procedure', 'Procedures'):
            return 'Procedure'
        if part in ('Specification', 'Specifications'):
            return 'Specification'
        if part in ('List', 'Datasheet'):
            return part
    return 'Unknown'

def detect_discipline_name(discipline: int | None) -> str:
    return DISCIPLINE_NAMES.get(discipline, 'Unknown')


# ---------------------------------------------------------------------------
# LibreOffice conversion
# ---------------------------------------------------------------------------

def convert_to_docx(doc_path: Path, out_dir: Path) -> Path | None:
    """Convert a .doc file to .docx using LibreOffice. Returns path to .docx or None on failure."""
    try:
        subprocess.run(
            [SOFFICE, '--headless', '--convert-to', 'docx', '--outdir', str(out_dir), str(doc_path)],
            capture_output=True, timeout=30
        )
        docx_path = out_dir / (doc_path.stem + '.docx')
        return docx_path if docx_path.exists() else None
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def is_revision_table(table) -> bool:
    """Detect revision history tables — first cell contains 'Rev' or 'Rev:'."""
    try:
        first_cell = table.rows[0].cells[0].text.strip().lower()
        return first_cell in ('rev:', 'rev', 'revision')
    except Exception:
        return False


def has_constraint(text: str) -> bool:
    return bool(CONSTRAINT_PATTERN.search(text))


def base_metadata(filename: str, file_path: Path) -> dict:
    discipline = detect_discipline(filename)
    return {
        'doc_id': detect_doc_id(filename),
        'doc_type': detect_doc_type(filename),
        'discipline': discipline,
        'discipline_name': detect_discipline_name(discipline),
        'source_folder': detect_source_folder(file_path),
        'revision': detect_revision(filename),
        'is_template': detect_is_template(filename),
    }


# ---------------------------------------------------------------------------
# PRC / SPC parser — section-heading chunking
# ---------------------------------------------------------------------------

def parse_prc(doc, meta: dict) -> list:
    """Chunk a PRC/SPC document by section headings."""
    chunks = []
    current_heading = None
    current_level = None
    current_body = []

    def flush(heading, level, body):
        text = '\n'.join(body).strip()
        if not text:
            return
        full_text = f'{heading}\n{text}' if heading else text
        chunks.append({
            'text': full_text,
            'metadata': {
                **meta,
                'section': heading or '',
                'section_level': level or 0,
                'is_constraint': has_constraint(full_text),
                'chunk_type': 'section',
            }
        })

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if not text:
            continue

        # Deleted / revision mark detection — skip struck-through paragraphs
        if all(run.font.strike for run in para.runs if run.text.strip()):
            continue

        if style in HEADING_STYLES:
            flush(current_heading, current_level, current_body)
            current_heading = text
            current_level = int(style[-1]) if style[-1].isdigit() else 1
            current_body = []
        else:
            current_body.append(text)

    flush(current_heading, current_level, current_body)
    return chunks


# ---------------------------------------------------------------------------
# LST .doc parser — table extraction + paragraph fallback
# ---------------------------------------------------------------------------

def table_to_chunks(table, meta: dict, sheet_name: str = '') -> list:
    """Convert a Word table to text chunks (same approach as XLS row_to_text)."""
    chunks = []
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

    if not rows:
        return chunks

    # Find header row — first row with 2+ non-empty, non-numeric cells
    header_idx = None
    for i, row in enumerate(rows):
        non_empty = [c for c in row if c and not c.replace('.', '').isdigit()]
        if len(non_empty) >= 2:
            header_idx = i
            break

    if header_idx is None:
        return chunks

    headers = rows[header_idx]

    for row_num, row in enumerate(rows[header_idx + 1:], start=header_idx + 2):
        if all(not c for c in row):
            continue

        row_dict = {
            headers[i]: row[i] if i < len(row) else ''
            for i in range(len(headers))
            if headers[i]
        }

        text = table_row_to_text(row_dict, meta['doc_id'], sheet_name)
        if not text.strip():
            continue

        empty = sum(1 for v in row_dict.values() if not v)
        needs_review = empty / max(len(row_dict), 1) > 0.5

        chunks.append({
            'text': text,
            'metadata': {
                **meta,
                'row': row_num,
                'chunk_type': 'table_row',
                'needs_llm_review': needs_review,
                'raw': row_dict,
            }
        })

    return chunks


def table_row_to_text(row: dict, doc_id: str, sheet_name: str) -> str:
    """Convert a table row to a natural language sentence."""

    def g(*keys):
        for k in keys:
            for rk, rv in row.items():
                if k.lower() in rk.lower() and rv and rv.strip():
                    return rv.strip()
        return ''

    # Control Logic (1-LST-0003): No. | ACTION | PERMISSIVE | INSTRUMENT NO. | Comments
    if '1-LST-0003' in doc_id:
        return (
            f"Control logic [{g('No.')}]: action [{g('ACTION')}], "
            f"permissive [{g('PERMISSIVE')}], "
            f"instrument [{g('INSTRUMENT')}], "
            f"comments [{g('Comments')}]."
        )

    # Electrical checklists (9-LST-07xx): Items to be checked
    if doc_id.startswith('9-LST-07'):
        return (
            f"Electrical system [{sheet_name}], "
            f"check item: [{g('Items to be checked', 'Item', 'Check', 'Description', 'Requirement')}]."
        )

    # Generic fallback
    parts = [f"{k}: {v}" for k, v in row.items() if v and v.strip()]
    return ' | '.join(parts)


def parse_lst_doc(doc, meta: dict) -> list:
    """Parse an LST .doc file: extract data tables, skip revision history."""
    chunks = []

    data_tables = [t for t in doc.tables if not is_revision_table(t)]

    for table in data_tables:
        chunks.extend(table_to_chunks(table, meta))

    # Fallback: if no tables or very few chunks, extract paragraphs
    if len(chunks) < 3:
        body = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if all(run.font.strike for run in para.runs if run.text.strip()):
                continue
            body.append(text)

        if body:
            full_text = '\n'.join(body)
            chunks.append({
                'text': full_text,
                'metadata': {
                    **meta,
                    'chunk_type': 'paragraphs',
                    'is_constraint': has_constraint(full_text),
                }
            })

    return chunks


# ---------------------------------------------------------------------------
# File-level parser
# ---------------------------------------------------------------------------

def parse_doc_file(file_path: Path, tmp_dir: Path) -> list:
    """Parse a single .doc or .docx file. Returns list of chunks."""
    filename = file_path.name

    if detect_is_hold(filename):
        return []

    meta = base_metadata(filename, file_path)

    # Convert .doc → .docx if needed
    if file_path.suffix.lower() == '.doc':
        docx_path = convert_to_docx(file_path, tmp_dir)
        if not docx_path:
            raise RuntimeError(f'LibreOffice conversion failed for {filename}')
    else:
        docx_path = file_path

    doc = docx.Document(str(docx_path))

    doc_type = meta['doc_type']
    if doc_type in ('PRC', 'SPC'):
        return parse_prc(doc, meta)
    else:
        return parse_lst_doc(doc, meta)


# ---------------------------------------------------------------------------
# Duplicate revision filter
# ---------------------------------------------------------------------------

def filter_highest_revision(files: list[Path]) -> list[Path]:
    """For files with the same base doc_id, keep only the highest revision."""
    groups: dict[str, Path] = {}
    ungrouped = []

    for f in files:
        doc_id = detect_doc_id(f.name)
        rev = detect_revision_number(f.name)
        if rev == -1:
            ungrouped.append(f)
            continue
        if doc_id not in groups or rev > detect_revision_number(groups[doc_id].name):
            groups[doc_id] = f

    return list(groups.values()) + ungrouped


# ---------------------------------------------------------------------------
# Batch runner
# ---------------------------------------------------------------------------

def parse_all_docs(root_dir: str | Path) -> tuple[list, list]:
    """
    Walk root_dir and parse all .doc/.docx files in List, Procedure,
    Datasheet, and Specification subfolders only.

    Returns (all_chunks, report).
    """
    root = Path(root_dir)
    all_chunks = []
    report = []

    doc_files = sorted(
        f for f in root.rglob('*')
        if f.suffix.lower() in ('.doc', '.docx')
        and any(part in ALLOWED_FOLDERS for part in f.parts)
    )

    doc_files = filter_highest_revision(doc_files)

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)

        for doc_file in doc_files:
            rel = doc_file.relative_to(root)
            print(f'Parsing: {rel}')
            try:
                chunks = parse_doc_file(doc_file, tmp)
                all_chunks.extend(chunks)
                report.append({
                    'file': str(rel),
                    'source_folder': detect_source_folder(doc_file),
                    'chunks': len(chunks),
                    'status': 'ok',
                })
                print(f'  → {len(chunks)} chunks')
            except Exception as e:
                print(f'  ERROR: {e}')
                report.append({'file': str(rel), 'status': 'error', 'error': str(e)})

    return all_chunks, report


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    import sys

    root = sys.argv[1] if len(sys.argv) > 1 else '.'

    chunks, report = parse_all_docs(root)

    out_dir = Path(__file__).parent

    # Merge with existing XLS chunks if present
    existing_path = out_dir / 'parsed_chunks.json'
    existing = []
    if existing_path.exists():
        with open(existing_path) as f:
            existing = json.load(f)
        print(f'\nLoaded {len(existing)} existing XLS chunks')

    all_chunks = existing + chunks

    with open(existing_path, 'w') as f:
        json.dump(all_chunks, f, indent=2, default=str)
    print(f'Saved {len(all_chunks)} total chunks → {existing_path}')

    # Save doc-only chunks separately for inspection
    doc_path = out_dir / 'parsed_chunks_doc.json'
    with open(doc_path, 'w') as f:
        json.dump(chunks, f, indent=2, default=str)
    print(f'Saved {len(chunks)} doc chunks → {doc_path}')

    # Summary report
    print('\n--- Parse Report ---')
    total = sum(r.get('chunks', 0) for r in report)
    errors = [r for r in report if r['status'] == 'error']

    for entry in report:
        if entry['status'] == 'error':
            print(f'  ERROR  {entry["file"]}: {entry["error"]}')
        else:
            print(f'  OK     {entry["file"]} — {entry["chunks"]} chunks')

    print(f'\nTotal: {total} doc chunks, {len(errors)} errors')
