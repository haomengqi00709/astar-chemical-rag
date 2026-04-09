"""
ingest.py — Parse engineering source documents into raw/ directory.

Routes each file to:
  Track A (narrative): Procedure, Spec text, DST templates
                       → raw/narrative/{discipline}/{doc_id}.md
  Track B (tables):    List rows, tabular Spec rows
                       → raw/tables/{doc_id}.csv

Routing rules:
  .doc/.docx:  PRC/SPC/DST → Track A  |  LST → Track B
  .xls/.xlsx:  LST/SPC     → Track B  |  DST → Track A
  .pdf:        always       → Track A

Usage:
    python src/ingest.py --source ../../drive-download-20260327T155115Z-1-001
    python src/ingest.py --source /path/to/docs --out ../raw
"""

import argparse
import csv
import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import openpyxl
import xlrd
import fitz          # PyMuPDF
import docx

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

from config import RAW_DIR, MANIFEST_PATH, TEMP_DIR as _CFG_TEMP

SOFFICE = next(
    (p for p in ['/opt/homebrew/bin/soffice', '/usr/bin/soffice', 'soffice']
     if shutil.which(p) or Path(p).exists()),
    'soffice'
)

TEMP_DIR = _CFG_TEMP
TEMP_DIR.mkdir(exist_ok=True)

DISCIPLINE_NAMES = {
    0: 'Administration',
    1: 'Process Technology',
    2: 'Civil & Structural',
    3: 'Structural',
    4: 'Equipment',
    5: 'Piping & Layout',
    6: 'Insulation',
    7: 'Coatings',
    8: 'Instrumentation & Control',
    9: 'Electrical',
}

SKIP_PATTERNS = re.compile(
    r'(HOLD|backup|Backup|~\$|\.DS_Store|\.xlk)', re.IGNORECASE
)


# ---------------------------------------------------------------------------
# Metadata detection
# ---------------------------------------------------------------------------

def detect_doc_id(filename: str) -> str:
    match = re.match(r'^(\d+-[A-Z]+-[\d\-]+)', filename)
    return match.group(1).rstrip('-') if match else filename.split('.')[0]

def detect_doc_type(filename: str) -> str:
    for t in ('LST', 'PRC', 'DST', 'SPC', 'CAL', 'FRM', 'REP', 'DAT', 'TAB'):
        if f'-{t}-' in filename.upper():
            return t
    return 'UNKNOWN'

def detect_discipline(filename: str) -> int | None:
    match = re.match(r'^(\d+)-', filename)
    return int(match.group(1)) if match else None

def detect_revision(filename: str) -> str:
    match = re.search(r'-R(\d+)', filename, re.IGNORECASE)
    return f'R{match.group(1)}' if match else 'R0'

def detect_source_folder(file_path: Path) -> str:
    for part in reversed(file_path.parts):
        if part in ('Procedure', 'Procedures'):
            return 'Procedure'
        if part in ('Specification', 'Specifications'):
            return 'Specification'
        if part == 'List':
            return 'List'
        if part == 'Datasheet':
            return 'Datasheet'
        if part == 'Calculation':
            return 'Calculation'
    return 'Unknown'

def detect_is_template(filename: str) -> bool:
    return bool(re.search(r'\[?[Tt]emplate\]?|\([Tt]emplate\)|XXXX', filename))

def get_meta(file_path: Path) -> dict:
    name = file_path.name
    disc = detect_discipline(name)
    return {
        'doc_id': detect_doc_id(name),
        'doc_type': detect_doc_type(name),
        'discipline': disc,
        'discipline_name': DISCIPLINE_NAMES.get(disc, 'Unknown') if disc is not None else 'Unknown',
        'source_folder': detect_source_folder(file_path),
        'revision': detect_revision(name),
        'is_template': detect_is_template(name),
        'source_file': str(file_path),
    }


# ---------------------------------------------------------------------------
# Track routing
# ---------------------------------------------------------------------------

def get_track(meta: dict, suffix: str) -> str:
    """
    Returns 'A' (narrative→wiki) or 'B' (tables→SQL).

    Rules:
      - .doc/.docx: LST → B, everything else → A
      - .xls/.xlsx: LST or SPC → B, DST or UNKNOWN → A
      - .pdf: always A
    """
    doc_type = meta['doc_type']
    if suffix in ('.doc', '.docx'):
        return 'B' if doc_type == 'LST' else 'A'
    if suffix in ('.xls', '.xlsx'):
        return 'B' if doc_type in ('LST', 'SPC') else 'A'
    return 'A'  # PDF


# ---------------------------------------------------------------------------
# Track A extractors → markdown
# ---------------------------------------------------------------------------

def _table_to_csv_rows(table) -> tuple[list[str], list[list[str]]]:
    """Extract headers and data rows from a python-docx table."""
    if not table.rows:
        return [], []
    headers = [c.text.strip() for c in table.rows[0].cells]
    rows = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if any(c for c in cells):
            rows.append(cells)
    return headers, rows


def extract_docx_as_markdown(path: Path,
                              doc_id: str | None = None,
                              tables_dir: Path | None = None) -> str:
    """
    Extract DOCX content as markdown, preserving heading structure.
    Tables are extracted in-line (at their correct position in the document).

    If doc_id and tables_dir are provided, inline tables are:
      - saved as CSVs to tables_dir/{doc_id}_table_{n}.csv
      - replaced in the markdown with [TABLE: {doc_id}_table_{n}]
    Otherwise tables are rendered as simple markdown tables inline.
    """
    from docx.oxml.ns import qn

    doc = docx.Document(str(path))
    lines = []
    table_counter = 0

    # Build a lookup from table XML element → table object
    table_map = {t._tbl: t for t in doc.tables}

    for child in doc.element.body:
        tag = child.tag

        if tag == qn('w:p'):
            # Paragraph
            para_text = child.text if hasattr(child, 'text') else ''
            # Get the actual paragraph text via python-docx
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            style = para.style.name if para.style else ''
            text = para.text.strip()
            if not text:
                continue
            if 'Heading 1' in style:
                lines.append(f'\n## {text}')
            elif 'Heading 2' in style:
                lines.append(f'\n### {text}')
            elif 'Heading 3' in style:
                lines.append(f'\n#### {text}')
            else:
                lines.append(text)

        elif tag == qn('w:tbl'):
            # Table — in document order
            table_counter += 1
            table = table_map.get(child)
            if table is None:
                continue

            headers, rows = _table_to_csv_rows(table)
            if not rows:
                continue

            if doc_id and tables_dir:
                # Save as CSV and insert placeholder
                slug = f'{doc_id}_table_{table_counter}'
                safe_slug = slug.replace('/', '_').replace('\\', '_')
                csv_path = tables_dir / f'{safe_slug}.csv'
                tables_dir.mkdir(parents=True, exist_ok=True)
                with open(csv_path, 'w', newline='', encoding='utf-8') as f:
                    import csv as _csv
                    writer = _csv.writer(f)
                    writer.writerow([f'# doc_id={doc_id} table={table_counter} inline=true'])
                    writer.writerow(headers)
                    writer.writerows(rows)
                lines.append(f'\n[TABLE: {safe_slug}]\n')
            else:
                # Render as markdown table inline
                if headers:
                    lines.append('\n| ' + ' | '.join(headers) + ' |')
                    lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
                for row in rows:
                    lines.append('| ' + ' | '.join(row) + ' |')
                lines.append('')

    return '\n'.join(lines)


def extract_doc_as_markdown(path: Path,
                             doc_id: str | None = None,
                             tables_dir: Path | None = None) -> str:
    """Convert .doc to .docx via LibreOffice, then extract."""
    tmp_dir = TEMP_DIR / f'doc_{path.stem[:20]}'
    tmp_dir.mkdir(exist_ok=True)
    try:
        result = subprocess.run(
            [SOFFICE, '--headless', '--convert-to', 'docx', '--outdir', str(tmp_dir), str(path)],
            capture_output=True, timeout=60
        )
        converted = tmp_dir / (path.stem + '.docx')
        if converted.exists():
            content = extract_docx_as_markdown(converted, doc_id=doc_id, tables_dir=tables_dir)
            converted.unlink()
            return content
        return f'[Could not convert {path.name}: {result.stderr.decode()[:200]}]'
    finally:
        import shutil as _shutil
        _shutil.rmtree(tmp_dir, ignore_errors=True)


def _iter_xls_rows(path: Path):
    """Yield (sheet_name, list_of_rows) for both .xls and .xlsx files."""
    suffix = path.suffix.lower()
    if suffix == '.xlsx':
        wb = openpyxl.load_workbook(path, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = [list(row) for row in ws.iter_rows(values_only=True)]
            yield sheet_name, rows
    else:  # .xls via xlrd
        wb = xlrd.open_workbook(str(path))
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            rows = [ws.row_values(i) for i in range(ws.nrows)]
            yield sheet_name, rows


def extract_xlsx_template_as_markdown(path: Path) -> str:
    """Extract DST/template XLS structure as markdown — describes fields, not data rows."""
    try:
        lines = []
        for sheet_name, rows in _iter_xls_rows(path):
            lines.append(f'\n## Sheet: {sheet_name}')
            row_count = 0
            for row in rows:
                cells = [str(c).strip() for c in row if c is not None and str(c).strip() not in ('None', '', 'none')]
                if cells:
                    lines.append('- ' + ' | '.join(cells))
                    row_count += 1
                if row_count > 80:
                    lines.append('- ...(truncated)')
                    break
        return '\n'.join(lines)
    except Exception as e:
        return f'[Could not read {path.name}: {e}]'


def extract_pdf_as_markdown(path: Path,
                             doc_id: str | None = None,
                             tables_dir: Path | None = None) -> str:
    """
    Extract PDF content as markdown, detecting tables per page.

    Uses pdfplumber for table detection (lattice + stream).
    Tables are extracted in reading order (top→bottom per page).

    If doc_id and tables_dir are provided, detected tables are:
      - saved as CSVs to tables_dir/{doc_id}_table_{n}.csv
      - replaced in the markdown with [TABLE: {doc_id}_table_{n}]
    Otherwise tables are rendered as markdown tables inline.

    Falls back to plain PyMuPDF text extraction if pdfplumber is unavailable.
    """
    try:
        import pdfplumber
    except ImportError:
        return _extract_pdf_pymupdf(path)

    import csv as _csv

    try:
        lines = []
        table_counter = 0

        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                lines.append(f'\n## Page {page_num}')

                # Detect tables on this page (bboxes in pdfplumber top-left coords)
                raw_tables = page.find_tables()

                if not raw_tables:
                    text = page.extract_text()
                    if text and text.strip():
                        lines.append(text.strip())
                    _append_pdf_annots(path, page_num, lines)
                    continue

                # Sort by vertical position (top of table, ascending = reading order)
                raw_tables_sorted = sorted(raw_tables, key=lambda t: t.bbox[1])

                prev_bottom = 0

                for table in raw_tables_sorted:
                    x0, top, x1, bottom = table.bbox

                    # --- text above this table ---
                    if top > prev_bottom:
                        region = page.within_bbox((0, prev_bottom, page.width, top))
                        text = region.extract_text()
                        if text and text.strip():
                            lines.append(text.strip())

                    # --- the table itself ---
                    data = table.extract()
                    if not data:
                        prev_bottom = bottom
                        continue

                    # First row = headers; remaining = data rows
                    raw_headers = data[0] if data else []
                    headers = [str(c).strip() if c else f'col_{i}'
                               for i, c in enumerate(raw_headers)]
                    data_rows = [
                        [str(c).strip() if c else '' for c in row]
                        for row in data[1:]
                        if any(c for c in row)
                    ]

                    if not data_rows:
                        prev_bottom = bottom
                        continue

                    table_counter += 1

                    if doc_id and tables_dir:
                        slug = f'{doc_id}_table_{table_counter}'
                        safe_slug = slug.replace('/', '_').replace('\\', '_')
                        csv_path = tables_dir / f'{safe_slug}.csv'
                        tables_dir.mkdir(parents=True, exist_ok=True)
                        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
                            writer = _csv.writer(f)
                            writer.writerow([
                                f'# doc_id={doc_id} table={table_counter}'
                                f' inline=true page={page_num}'
                            ])
                            writer.writerow(headers)
                            writer.writerows(data_rows)
                        lines.append(f'\n[TABLE: {safe_slug}]\n')
                    else:
                        lines.append('\n| ' + ' | '.join(headers) + ' |')
                        lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
                        for row in data_rows:
                            lines.append('| ' + ' | '.join(row) + ' |')
                        lines.append('')

                    prev_bottom = bottom

                # --- text below last table ---
                if prev_bottom < page.height:
                    region = page.within_bbox((0, prev_bottom, page.width, page.height))
                    text = region.extract_text()
                    if text and text.strip():
                        lines.append(text.strip())

                _append_pdf_annots(path, page_num, lines)

        return '\n'.join(lines)

    except Exception as e:
        return f'[Could not read {path.name}: {e}]'


def _extract_pdf_pymupdf(path: Path) -> str:
    """Fallback: plain text extraction via PyMuPDF (no table detection)."""
    try:
        doc = fitz.open(str(path))
        lines = []
        for page_num, page in enumerate(doc, 1):
            text = page.get_text('text').strip()
            if text:
                lines.append(f'\n## Page {page_num}\n{text}')
            _append_pdf_annots(path, page_num, lines)
        return '\n'.join(lines)
    except Exception as e:
        return f'[Could not read {path.name}: {e}]'


def _append_pdf_annots(path: Path, page_num: int, lines: list) -> None:
    """Append any human annotations from a PyMuPDF page to lines list."""
    try:
        doc = fitz.open(str(path))
        page = doc[page_num - 1]
        for annot in page.annots():
            content = annot.info.get('content', '').strip()
            if content:
                lines.append(f'\n> [Annotation p.{page_num}] {content}')
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Track B extractors → CSV
# ---------------------------------------------------------------------------

def _find_header_row(ws) -> int:
    """Find the first row with 3+ non-empty cells — likely the header."""
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        non_empty = [c for c in row if c is not None and str(c).strip()]
        if len(non_empty) >= 3:
            return i
    return 0


def extract_xlsx_as_csv_text(path: Path) -> list[dict]:
    """
    Extract XLS/XLSX tabular data.
    Returns list of {sheet, headers, rows} dicts — one per sheet.
    """
    results = []
    try:
        for sheet_name, all_rows in _iter_xls_rows(path):
            if not all_rows:
                continue

            # Find header row: first row with 3+ non-empty cells
            header_idx = 0
            for i, row in enumerate(all_rows):
                non_empty = [c for c in row if c is not None and str(c).strip()]
                if len(non_empty) >= 3:
                    header_idx = i
                    break

            raw_headers = all_rows[header_idx]
            headers = [str(h).strip() if h is not None else f'col_{i}'
                       for i, h in enumerate(raw_headers)]
            # Deduplicate empty headers
            seen = {}
            clean_headers = []
            for h in headers:
                if not h or h == 'None':
                    h = f'col_{len(clean_headers)}'
                if h in seen:
                    seen[h] += 1
                    h = f'{h}_{seen[h]}'
                else:
                    seen[h] = 0
                clean_headers.append(h)

            data_rows = []
            for row in all_rows[header_idx + 1:]:
                cells = [str(c).strip() if c is not None else '' for c in row]
                # Skip completely empty rows
                if any(c for c in cells):
                    data_rows.append(cells)

            if data_rows:
                results.append({
                    'sheet': sheet_name,
                    'headers': clean_headers,
                    'rows': data_rows,
                })
    except Exception as e:
        print(f'  [WARN] Could not parse {path.name}: {e}')
    return results


def extract_docx_table_as_csv(path: Path) -> list[dict]:
    """Extract Word document tables as CSV-compatible structure."""
    results = []
    try:
        doc = docx.Document(str(path))
        for i, table in enumerate(doc.tables):
            if not table.rows:
                continue
            headers = [c.text.strip() for c in table.rows[0].cells]
            rows = []
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                results.append({
                    'sheet': f'table_{i+1}',
                    'headers': headers,
                    'rows': rows,
                })
    except Exception as e:
        print(f'  [WARN] Could not parse {path.name}: {e}')
    return results


# ---------------------------------------------------------------------------
# Writers
# ---------------------------------------------------------------------------

def write_narrative(meta: dict, content: str, out_dir: Path) -> Path:
    """Write Track A markdown to raw/narrative/{discipline}/{doc_id}.md"""
    disc_str = f"{meta['discipline']} {meta['discipline_name']}" if meta['discipline'] is not None else 'Unknown'
    folder = out_dir / 'narrative' / disc_str
    folder.mkdir(parents=True, exist_ok=True)

    header = f"""---
doc_id: {meta['doc_id']}
doc_type: {meta['doc_type']}
discipline: {meta['discipline']}
discipline_name: {meta['discipline_name']}
source_folder: {meta['source_folder']}
revision: {meta['revision']}
is_template: {meta['is_template']}
track: A
---

# {meta['doc_id']} — {meta['source_folder']}

"""
    out_path = folder / f"{meta['doc_id']}.md"
    out_path.write_text(header + content, encoding='utf-8')
    return out_path


def write_tables(meta: dict, sheets: list[dict], out_dir: Path) -> list[Path]:
    """Write Track B CSVs to raw/tables/{doc_id}_{sheet}.csv"""
    folder = out_dir / 'tables'
    folder.mkdir(parents=True, exist_ok=True)
    out_paths = []

    for sheet_data in sheets:
        sheet_suffix = f"_{sheet_data['sheet']}" if len(sheets) > 1 else ''
        safe_doc_id = meta['doc_id'].replace('/', '_').replace('\\', '_')
        out_path = folder / f"{safe_doc_id}{sheet_suffix}.csv"

        with open(out_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            # Write metadata row as comment
            writer.writerow([f'# doc_id={meta["doc_id"]} discipline={meta["discipline"]} revision={meta["revision"]}'])
            writer.writerow(sheet_data['headers'])
            writer.writerows(sheet_data['rows'])

        out_paths.append(out_path)
    return out_paths


# ---------------------------------------------------------------------------
# Main processing
# ---------------------------------------------------------------------------

def process_file(file_path: Path, raw_dir: Path) -> dict | None:
    """Process one file. Returns manifest entry or None if skipped."""
    name = file_path.name
    suffix = file_path.suffix.lower()

    if SKIP_PATTERNS.search(name):
        return None
    if suffix not in ('.doc', '.docx', '.xls', '.xlsx', '.pdf'):
        return None

    meta = get_meta(file_path)
    track = get_track(meta, suffix)
    meta['track'] = track

    print(f'  [{track}] {meta["doc_id"]:30s} {meta["doc_type"]:6s} {suffix}')

    try:
        if track == 'A':
            tables_dir = raw_dir / 'tables'
            doc_id = meta['doc_id']
            if suffix == '.docx':
                content = extract_docx_as_markdown(file_path, doc_id=doc_id, tables_dir=tables_dir)
            elif suffix == '.doc':
                content = extract_doc_as_markdown(file_path, doc_id=doc_id, tables_dir=tables_dir)
            elif suffix in ('.xls', '.xlsx'):
                content = extract_xlsx_template_as_markdown(file_path)
            else:  # pdf
                content = extract_pdf_as_markdown(file_path, doc_id=doc_id, tables_dir=tables_dir)

            out = write_narrative(meta, content, raw_dir)
            meta['output'] = str(out)
            # Count inline tables extracted
            inline_tables = content.count('[TABLE:')
            if inline_tables:
                meta['inline_tables'] = inline_tables
                print(f'    → {inline_tables} inline table(s) extracted to raw/tables/')

        else:  # Track B
            if suffix in ('.xls', '.xlsx'):
                sheets = extract_xlsx_as_csv_text(file_path)
            else:  # .doc/.docx with tables
                if suffix == '.doc':
                    tmp_dir = TEMP_DIR / f'b_{file_path.stem[:20]}'
                    tmp_dir.mkdir(exist_ok=True)
                    try:
                        subprocess.run(
                            [SOFFICE, '--headless', '--convert-to', 'docx', '--outdir', str(tmp_dir), str(file_path)],
                            capture_output=True, timeout=60
                        )
                        converted = tmp_dir / (file_path.stem + '.docx')
                        sheets = extract_docx_table_as_csv(converted) if converted.exists() else []
                    finally:
                        import shutil as _shutil
                        _shutil.rmtree(tmp_dir, ignore_errors=True)
                else:
                    sheets = extract_docx_table_as_csv(file_path)

            if sheets:
                out_paths = write_tables(meta, sheets, raw_dir)
                meta['output'] = [str(p) for p in out_paths]
                meta['row_count'] = sum(len(s['rows']) for s in sheets)
            else:
                print(f'    [WARN] No table data found — skipping Track B for {name}')
                return None

    except Exception as e:
        print(f'    [ERROR] {name}: {e}')
        meta['error'] = str(e)

    return meta


def ingest(source_dir: Path, raw_dir: Path) -> list[dict]:
    """Walk source_dir, process all files, return manifest."""
    all_files = sorted(source_dir.rglob('*'))
    eligible = [f for f in all_files if f.is_file() and not SKIP_PATTERNS.search(f.name)]

    print(f'Found {len(eligible)} files in {source_dir}')
    print()

    manifest = []
    track_a, track_b, skipped = 0, 0, 0

    for f in eligible:
        entry = process_file(f, raw_dir)
        if entry:
            manifest.append(entry)
            if entry.get('track') == 'A':
                track_a += 1
            else:
                track_b += 1
        else:
            skipped += 1

    print(f'\nDone: {track_a} Track A (narrative), {track_b} Track B (tables), {skipped} skipped')
    return manifest


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description='Ingest engineering documents into raw/')
    parser.add_argument('--source', required=True, help='Root directory of source documents')
    parser.add_argument('--out', default=str(RAW_DIR), help='Output raw/ directory')
    args = parser.parse_args()

    source = Path(args.source).resolve()
    raw_dir = Path(args.out).resolve()
    raw_dir.mkdir(parents=True, exist_ok=True)

    print(f'Source: {source}')
    print(f'Output: {raw_dir}')
    print()

    manifest = ingest(source, raw_dir)

    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, default=str))
    print(f'Manifest saved → {MANIFEST_PATH}')


if __name__ == '__main__':
    main()
